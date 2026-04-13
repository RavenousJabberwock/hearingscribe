"""
HearingScribe — transcription backend
Runs as a subprocess spawned by Electron main.js.
Progress is written to stdout as JSON lines.
Errors go to stderr.
"""

import sys, os, json, argparse, time, warnings
warnings.filterwarnings("ignore")

# ── Load model paths written by setup.js ──────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PATHS_FILE = os.path.join(SCRIPT_DIR, 'model_paths.json')

if not os.path.exists(PATHS_FILE):
    print("FATAL: model_paths.json not found. Please re-run setup.js.", file=sys.stderr)
    sys.exit(1)

with open(PATHS_FILE) as f:
    MODEL_PATHS = json.load(f)

WHISPER_DIR    = MODEL_PATHS['whisper_dir']
PYANNOTE_CACHE = MODEL_PATHS['pyannote_cache']
PYANNOTE_FLAT  = MODEL_PATHS.get('pyannote_flat', PYANNOTE_CACHE)

# Force HuggingFace to use our local cache only — no internet calls ever
# Use the flat dir if it exists (packaged build), otherwise fall back to HF cache (dev)
if os.path.isdir(PYANNOTE_FLAT) and os.listdir(PYANNOTE_FLAT):
    os.environ['HF_HOME'] = PYANNOTE_FLAT
else:
    os.environ['HF_HOME'] = PYANNOTE_CACHE
os.environ['TRANSFORMERS_OFFLINE'] = '1'
os.environ['HF_DATASETS_OFFLINE']  = '1'

# ── helpers ───────────────────────────────────────────────────────────────────
def emit(data):
    print(json.dumps(data), flush=True)

def progress(step=None, pct=None, label=None, log=None,
             log_type='info', speakers=None, remain=None):
    msg = {}
    if step     is not None: msg['step']    = step
    if pct      is not None: msg['pct']     = pct
    if label    is not None: msg['label']   = label
    if log      is not None: msg['log']     = log
    if log_type != 'info':   msg['logType'] = log_type
    if speakers is not None: msg['speakers']= speakers
    if remain   is not None: msg['remain']  = remain
    emit(msg)

# ── args ──────────────────────────────────────────────────────────────────────
parser = argparse.ArgumentParser()
parser.add_argument('--input',  required=True)
parser.add_argument('--output', required=True)
parser.add_argument('--format', choices=['txt','docx'], default='docx')
args = parser.parse_args()

start_time = time.time()

try:
    # ── Step 0: load audio ────────────────────────────────────────────────────
    progress(step=0, pct=3, label='Loading audio file…',
             log='Reading: ' + os.path.basename(args.input))

    import torch, torchaudio

    waveform, sample_rate = torchaudio.load(args.input)
    duration_sec = waveform.shape[1] / sample_rate

    if sample_rate != 16000:
        waveform = torchaudio.transforms.Resample(sample_rate, 16000)(waveform)
    if waveform.shape[0] > 1:
        waveform = waveform.mean(dim=0, keepdim=True)

    audio_np = waveform.squeeze().numpy()
    progress(pct=8,
             log=f'Audio loaded — {duration_sec/60:.1f} min',
             log_type='ok',
             remain=duration_sec * 0.45)

    # ── Step 1: transcribe ────────────────────────────────────────────────────
    progress(step=1, pct=10, label='Transcribing speech…',
             log='Loading Whisper large-v3')

    import whisper
    model = whisper.load_model("large-v3", download_root=WHISPER_DIR)
    progress(pct=14, log='Whisper ready', log_type='ok')
    progress(pct=16, log='Running transcription — longest step')

    result   = model.transcribe(audio_np, language='en', word_timestamps=True, verbose=False)
    segments = result['segments']
    progress(pct=55, log=f'Transcription complete — {len(segments)} segments', log_type='ok')

    # ── Step 2: diarize ───────────────────────────────────────────────────────
    progress(step=2, pct=57, label='Identifying speakers…',
             log='Loading pyannote from local cache')

    from pyannote.audio import Pipeline
    pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-community-1")
    device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
    pipeline = pipeline.to(device)
    progress(pct=60, log=f'Diarization ready · device={device}', log_type='ok')

    import tempfile, soundfile as sf
    with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as tmp:
        tmp_path = tmp.name
    sf.write(tmp_path, audio_np, 16000)
    diarization = pipeline(tmp_path)
    os.unlink(tmp_path)

    turns       = [{'start': t.start, 'end': t.end, 'speaker': s}
                   for t, _, s in diarization.itertracks(yield_label=True)]
    speaker_ids = sorted(set(t['speaker'] for t in turns))
    label_map   = {sid: f'Speaker{i+1}' for i, sid in enumerate(speaker_ids)}
    num_speakers = len(speaker_ids)

    progress(pct=80, log=f'Detected {num_speakers} speaker(s)', log_type='ok',
             speakers=str(num_speakers))

    # ── Merge ─────────────────────────────────────────────────────────────────
    def speaker_at(t0, t1):
        best, best_ov = 'Speaker1', 0.0
        for t in turns:
            ov = max(0, min(t1, t['end']) - max(t0, t['start']))
            if ov > best_ov:
                best_ov = ov
                best    = label_map[t['speaker']]
        return best

    out_lines = []
    cur_spk = cur_txt = cur_start = None
    for seg in segments:
        spk = speaker_at(seg['start'], seg['end'])
        if spk != cur_spk:
            if cur_spk:
                mm, ss = divmod(int(cur_start), 60)
                out_lines.append({'speaker': cur_spk, 'time': f'{mm:02d}:{ss:02d}',
                                   'text': ' '.join(cur_txt).strip()})
            cur_spk, cur_txt, cur_start = spk, [seg['text'].strip()], seg['start']
        else:
            cur_txt.append(seg['text'].strip())
    if cur_spk:
        mm, ss = divmod(int(cur_start), 60)
        out_lines.append({'speaker': cur_spk, 'time': f'{mm:02d}:{ss:02d}',
                           'text': ' '.join(cur_txt).strip()})

    progress(pct=88, log=f'Merged into {len(out_lines)} speaker turns', log_type='ok')

    # ── Step 3: write output ──────────────────────────────────────────────────
    progress(step=3, pct=90, label='Generating transcript document…',
             log=f'Writing {args.format.upper()}')

    ts       = time.strftime('%Y-%m-%d %H:%M')
    src_name = os.path.basename(args.input)

    if args.format == 'txt':
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(f'HEARING TRANSCRIPT\nGenerated: {ts}\nSource: {src_name}\nSpeakers: {num_speakers}\n{"="*60}\n\n')
            for line in out_lines:
                f.write(f'[{line["speaker"]}]  {line["time"]}\n{line["text"]}\n\n')
    else:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for sec in doc.sections:
            sec.top_margin = sec.bottom_margin = Inches(1)
            sec.left_margin = sec.right_margin = Inches(1.25)

        h = doc.add_heading('HEARING TRANSCRIPT', level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, txt in enumerate([f'Generated: {ts}', f'Source: {src_name}', f'Speakers detected: {num_speakers}']):
            r = meta.add_run(('\n' if i else '') + txt)
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        doc.add_paragraph()

        colors = [RGBColor(0x1a,0x5c,0x8c), RGBColor(0x8c,0x2a,0x1a),
                  RGBColor(0x1a,0x7a,0x45), RGBColor(0x6a,0x1a,0x8c)]
        def spk_color(lbl):
            return colors[(int(lbl.replace('Speaker','') or 1)-1) % len(colors)]

        for line in out_lines:
            p    = doc.add_paragraph()
            head = p.add_run(f'[{line["speaker"]}]  {line["time"]}')
            head.bold = True; head.font.size = Pt(9); head.font.color.rgb = spk_color(line['speaker'])
            p.add_run('\n' + line['text']).font.size = Pt(11)
            p.paragraph_format.space_after = Pt(10)

        doc.save(args.output)

    elapsed = time.time() - start_time
    progress(pct=100, label='Complete!',
             log=f'Done in {elapsed:.0f}s — {os.path.basename(args.output)}',
             log_type='ok')

except Exception as e:
    import traceback
    print(traceback.format_exc(), file=sys.stderr, flush=True)
    sys.exit(1)
